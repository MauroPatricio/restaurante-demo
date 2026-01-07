"""
Script para gerar documentação técnica do Sistema de Gestão de Mesas
Bilíngue: Português e Inglês
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def add_heading_custom(doc, text, level=1):
    """Adiciona título com formatação customizada"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code):
    """Adiciona bloco de código"""
    para = doc.add_paragraph(code)
    para.style = 'Intense Quote'
    return para

def create_technical_document():
    """Cria documento técnico completo"""
    doc = Document()
    
    # ======================================
    # CAPA / COVER PAGE
    # ======================================
    title = doc.add_heading('Sistema de Gestão de Estado de Mesas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Table State Management System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()
    
    info = doc.add_paragraph('Documentação Técnica | Technical Documentation')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    version = doc.add_paragraph('Versão 1.0 | Version 1.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph('Dezembro 2025 | December 2025')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ======================================
    # 1. INTRODUÇÃO / INTRODUCTION
    # ======================================
    add_heading_custom(doc, '1. Introdução / Introduction', 1)
    
    add_heading_custom(doc, '1.1 Visão Geral / Overview', 2)
    
    doc.add_paragraph(
        'PT: O Sistema de Gestão de Estado de Mesas é uma solução completa para controlar '
        'o ciclo de vida das mesas em um restaurante, incluindo transições automáticas de '
        'estado, associação de pedidos a sessões, e controles de autorização para liberação '
        'de mesas.'
    )
    
    doc.add_paragraph(
        'EN: The Table State Management System is a comprehensive solution to control the '
        'lifecycle of tables in a restaurant, including automatic state transitions, order '
        'association with sessions, and authorization controls for table release.'
    )
    
    add_heading_custom(doc, '1.2 Objetivos / Objectives', 2)
    
    doc.add_paragraph('PT: Principais objetivos do sistema:')
    doc.add_paragraph('EN: Main system objectives:', style='List Bullet')
    
    objectives_pt = [
        'Automatizar transição de mesa livre para ocupada ao criar pedido',
        'Associar todos os pedidos a sessões de ocupação',
        'Permitir apenas managers e waiters liberarem mesas',
        'Manter histórico completo de sessões',
        'Rastrear receita por sessão de mesa',
        'Auditoria de mudanças de estado'
    ]
    
    objectives_en = [
        'Automate table transition from free to occupied when creating order',
        'Associate all orders with occupation sessions',
        'Allow only managers and waiters to free tables',
        'Maintain complete session history',
        'Track revenue per table session',
        'Audit state changes'
    ]
    
    for pt, en in zip(objectives_pt, objectives_en):
        doc.add_paragraph(f'🇵🇹 {pt}', style='List Bullet')
        doc.add_paragraph(f'🇬🇧 {en}', style='List Bullet')
    
    doc.add_page_break()
    
    # ======================================
    # 2. ARQUITETURA / ARCHITECTURE
    # ======================================
    add_heading_custom(doc, '2. Arquitetura do Sistema / System Architecture', 1)
    
    add_heading_custom(doc, '2.1 Componentes Principais / Main Components', 2)
    
    doc.add_paragraph('PT: Backend (Node.js + Express + MongoDB)')
    doc.add_paragraph('EN: Backend (Node.js + Express + MongoDB)')
    
    components_pt = [
        'Modelos de Dados: TableSession, Table (atualizado), Order (atualizado)',
        'Controladores: tableStateController.js',
        'Middleware: tableValidation.js, canFreeTable',
        'Rotas: Endpoints REST para gerenciamento de sessões'
    ]
    
    components_en = [
        'Data Models: TableSession, Table (updated), Order (updated)',
        'Controllers: tableStateController.js',
        'Middleware: tableValidation.js, canFreeTable',
        'Routes: REST endpoints for session management'
    ]
    
    for pt, en in zip(components_pt, components_en):
        doc.add_paragraph(f'🇵🇹 {pt}', style='List Bullet')
        doc.add_paragraph(f'🇬🇧 {en}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('PT: Frontend Admin Dashboard (React + Vite)')
    doc.add_paragraph('EN: Frontend Admin Dashboard (React + Vite)')
    
    frontend_pt = [
        'TableSessionModal: Modal de visualização de sessão',
        'Tables Page: Página de gerenciamento de mesas',
        'API Service: Funções de comunicação com backend',
        'i18n: Suporte a 4 idiomas (PT, EN, ES, FR)'
    ]
    
    frontend_en = [
        'TableSessionModal: Session visualization modal',
        'Tables Page: Table management page',
        'API Service: Backend communication functions',
        'i18n: Support for 4 languages (PT, EN, ES, FR)'
    ]
    
    for pt, en in zip(frontend_pt, frontend_en):
        doc.add_paragraph(f'🇵🇹 {pt}', style='List Bullet')
        doc.add_paragraph(f'🇬🇧 {en}', style='List Bullet')
    
    doc.add_page_break()
    
    # ======================================
    # 3. MODELOS DE DADOS / DATA MODELS
    # ======================================
    add_heading_custom(doc, '3. Modelos de Dados / Data Models', 1)
    
    add_heading_custom(doc, '3.1 TableSession', 2)
    
    doc.add_paragraph('PT: Modelo para rastrear ciclos de ocupação de mesas.')
    doc.add_paragraph('EN: Model to track table occupation cycles.')
    
    add_code_block(doc, '''
{
  table: ObjectId (ref: Table),
  restaurant: ObjectId (ref: Restaurant),
  startedAt: Date,
  endedAt: Date,
  startedBy: ObjectId (ref: User),
  endedBy: ObjectId (ref: User),
  status: 'active' | 'closed',
  totalRevenue: Number,
  orderCount: Number
}
    ''')
    
    doc.add_paragraph('PT: Campos Principais / EN: Main Fields:', style='Heading 3')
    
    fields = [
        ('table', 'Referência para a mesa / Table reference'),
        ('restaurant', 'Referência para o restaurante / Restaurant reference'),
        ('startedAt', 'Data/hora de início da sessão / Session start date/time'),
        ('endedAt', 'Data/hora de término da sessão / Session end date/time'),
        ('status', 'Status da sessão (active/closed) / Session status'),
        ('totalRevenue', 'Receita total da sessão / Total session revenue'),
        ('orderCount', 'Número de pedidos na sessão / Number of orders in session')
    ]
    
    for field, desc in fields:
        doc.add_paragraph(f'• {field}: {desc}', style='List Bullet')
    
    add_heading_custom(doc, '3.2 Table (Atualizado / Updated)', 2)
    
    doc.add_paragraph('PT: Campos adicionados ao modelo Table existente:')
    doc.add_paragraph('EN: Fields added to existing Table model:')
    
    add_code_block(doc, '''
{
  currentSessionId: ObjectId (ref: TableSession),
  lastStatusChange: Date,
  statusChangedBy: ObjectId (ref: User)
}
    ''')
    
    add_heading_custom(doc, '3.3 Order (Atualizado / Updated)', 2)
    
    doc.add_paragraph('PT: Campo adicionado ao modelo Order existente:')
    doc.add_paragraph('EN: Field added to existing Order model:')
    
    add_code_block(doc, '''
{
  tableSession: ObjectId (ref: TableSession)
}
    ''')
    
    doc.add_page_break()
    
    # ======================================
    # 4. API ENDPOINTS
    # ======================================
    add_heading_custom(doc, '4. API Endpoints', 1)
    
    add_heading_custom(doc, '4.1 GET /api/tables/:id/current-session', 2)
    
    doc.add_paragraph('PT: Obtém a sessão atual de uma mesa com todos os pedidos.')
    doc.add_paragraph('EN: Gets current session of a table with all orders.')
    
    doc.add_paragraph('Autenticação / Authentication: Bearer Token', style='Heading 3')
    doc.add_paragraph('Autorização / Authorization: Qualquer usuário autenticado / Any authenticated user')
    
    doc.add_paragraph('Resposta / Response:', style='Heading 3')
    add_code_block(doc, '''
{
  "table": { "_id", "number", "status", ... },
  "session": { "startedAt", "status", ... },
  "orders": [ { "items", "total", ... } ],
  "stats": {
    "orderCount": 3,
    "totalRevenue": 1500,
    "sessionDuration": 45
  }
}
    ''')
    
    add_heading_custom(doc, '4.2 POST /api/tables/:id/free', 2)
    
    doc.add_paragraph('PT: Libera uma mesa ocupada, encerrando a sessão atual.')
    doc.add_paragraph('EN: Frees an occupied table, closing current session.')
    
    doc.add_paragraph('Autenticação / Authentication: Bearer Token', style='Heading 3')
    doc.add_paragraph('Autorização / Authorization: manager, waiter, owner')
    
    doc.add_paragraph('Ações / Actions:', style='Heading 3')
    actions = [
        'Fecha a sessão ativa / Closes active session',
        'Calcula receita total / Calculates total revenue',
        'Atualiza status da mesa para "free" / Updates table status to "free"',
        'Registra quem liberou / Records who freed the table'
    ]
    for action in actions:
        doc.add_paragraph(f'• {action}', style='List Bullet')
    
    add_heading_custom(doc, '4.3 GET /api/tables/:id/session-history', 2)
    
    doc.add_paragraph('PT: Retorna histórico de sessões passadas de uma mesa.')
    doc.add_paragraph('EN: Returns history of past sessions for a table.')
    
    doc.add_paragraph('Autorização / Authorization: manager, waiter, owner')
    
    doc.add_paragraph('Parâmetros / Parameters:', style='Heading 3')
    doc.add_paragraph('• limit: Número de sessões (padrão: 10) / Number of sessions (default: 10)')
    doc.add_paragraph('• page: Página de resultados (padrão: 1) / Results page (default: 1)')
    
    doc.add_page_break()
    
    # ======================================
    # 5. MIDDLEWARE E VALIDAÇÃO
    # ======================================
    add_heading_custom(doc, '5. Middleware e Validação / Middleware and Validation', 1)
    
    add_heading_custom(doc, '5.1 validateAndOccupyTable', 2)
    
    doc.add_paragraph('PT: Middleware que valida o status da mesa antes de criar pedido.')
    doc.add_paragraph('EN: Middleware that validates table status before creating order.')
    
    doc.add_paragraph('Funcionalidades / Features:', style='Heading 3')
    features = [
        'Valida se mesa não está fechada ou em limpeza / Validates table is not closed or cleaning',
        'Auto-transição free → occupied ao criar pedido / Auto-transition free → occupied when creating order',
        'Cria nova TableSession automaticamente / Creates new TableSession automatically',
        'Vincula pedido à sessão / Links order to session'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    add_heading_custom(doc, '5.2 canFreeTable', 2)
    
    doc.add_paragraph('PT: Middleware de autorização para liberação de mesas.')
    doc.add_paragraph('EN: Authorization middleware for table release.')
    
    doc.add_paragraph('Regra / Rule:', style='Heading 3')
    doc.add_paragraph('• Apenas roles: manager, waiter, owner / Only roles: manager, waiter, owner')
    doc.add_paragraph('• Retorna 403 Forbidden para outros usuários / Returns 403 Forbidden for other users')
    
    doc.add_page_break()
    
    # ======================================
    # 6. FRONTEND - ADMIN DASHBOARD
    # ======================================
    add_heading_custom(doc, '6. Frontend - Admin Dashboard', 1)
    
    add_heading_custom(doc, '6.1 TableSessionModal Component', 2)
    
    doc.add_paragraph('PT: Modal para visualização de sessão de mesa com estatísticas e pedidos.')
    doc.add_paragraph('EN: Modal for viewing table session with statistics and orders.')
    
    doc.add_paragraph('Funcionalidades / Features:', style='Heading 3')
    modal_features = [
        'Exibe informações da mesa / Displays table information',
        'Mostra estatísticas da sessão (duração, pedidos, receita) / Shows session stats',
        'Lista todos os pedidos da sessão / Lists all session orders',
        'Botão "Liberar Mesa" (manager/waiter) / "Free Table" button',
        'Confirmação antes de liberar / Confirmation before freeing',
        'Design responsivo / Responsive design'
    ]
    for feature in modal_features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    add_heading_custom(doc, '6.2 Tables Page (Atualizado / Updated)', 2)
    
    doc.add_paragraph('PT: Página de gerenciamento de mesas com nova funcionalidade.')
    doc.add_paragraph('EN: Table management page with new functionality.')
    
    doc.add_paragraph('Mudanças / Changes:', style='Heading 3')
    changes = [
        'Botão "👁️ Ver Pedidos" em cada mesa / "👁️ View Orders" button on each table',
        'Função handleViewSession() / handleViewSession() function',
        'Função handleFreeTable() / handleFreeTable() function',
        'Integração com TableSessionModal / Integration with TableSessionModal',
        'Refresh automático após liberar / Auto-refresh after freeing'
    ]
    for change in changes:
        doc.add_paragraph(f'• {change}', style='List Bullet')
    
    doc.add_page_break()
    
    # ======================================
    # 7. FLUXOS DE TRABALHO
    # ======================================
    add_heading_custom(doc, '7. Fluxos de Trabalho / Workflows', 1)
    
    add_heading_custom(doc, '7.1 Cliente Cria Primeiro Pedido / Client Creates First Order', 2)
    
    doc.add_paragraph('PT: Fluxo automático de transição free → occupied')
    doc.add_paragraph('EN: Automatic transition flow free → occupied')
    
    steps_pt = [
        '1. Cliente escaneia QR Code',
        '2. Adiciona itens ao carrinho',
        '3. Clica em "Fazer Pedido"',
        '4. Backend valida status da mesa',
        '5. Se mesa = "free": cria TableSession',
        '6. Mesa.status = "occupied"',
        '7. Cria Order vinculado à sessão',
        '8. Cliente vê confirmação'
    ]
    
    steps_en = [
        '1. Client scans QR Code',
        '2. Adds items to cart',
        '3. Clicks "Place Order"',
        '4. Backend validates table status',
        '5. If table = "free": creates TableSession',
        '6. Table.status = "occupied"',
        '7. Creates Order linked to session',
        '8. Client sees confirmation'
    ]
    
    for pt, en in zip(steps_pt, steps_en):
        doc.add_paragraph(f'🇵🇹 {pt}')
        doc.add_paragraph(f'🇬🇧 {en}')
    
    add_heading_custom(doc, '7.2 Manager/Waiter Libera Mesa / Frees Table', 2)
    
    steps2_pt = [
        '1. Manager acessa Admin Dashboard',
        '2. Clica em "👁️" na mesa ocupada',
        '3. Vê estatísticas e pedidos',
        '4. Clica em "Liberar Mesa"',
        '5. Confirma ação',
        '6. Backend fecha sessão',
        '7. Mesa.status = "free"',
        '8. UI atualiza automaticamente'
    ]
    
    steps2_en = [
        '1. Manager accesses Admin Dashboard',
        '2. Clicks "👁️" on occupied table',
        '3. Views statistics and orders',
        '4. Clicks "Free Table"',
        '5. Confirms action',
        '6. Backend closes session',
        '7. Table.status = "free"',
        '8. UI updates automatically'
    ]
    
    for pt, en in zip(steps2_pt, steps2_en):
        doc.add_paragraph(f'🇵🇹 {pt}')
        doc.add_paragraph(f'🇬🇧 {en}')
    
    doc.add_page_break()
    
    # ======================================
    # 8. SEGURANÇA E AUTORIZAÇÃO
    # ======================================
    add_heading_custom(doc, '8. Segurança e Autorização / Security and Authorization', 1)
    
    add_heading_custom(doc, '8.1 Controle de Acesso / Access Control', 2)
    
    doc.add_paragraph('PT: Matriz de permissões por role:')
    doc.add_paragraph('EN: Permission matrix by role:')
    
    # Criar tabela
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header
    headers = ['Ação / Action', 'Client', 'Waiter', 'Manager', 'Owner']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Data
    permissions = [
        ['Ver sessão / View session', '❌', '✅', '✅', '✅'],
        ['Liberar mesa / Free table', '❌', '✅', '✅', '✅'],
        ['Ver histórico / View history', '❌', '✅', '✅', '✅'],
        ['Criar pedido / Create order', '✅', '✅', '✅', '✅']
    ]
    
    for i, row_data in enumerate(permissions):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '8.2 Middleware de Autorização / Authorization Middleware', 2)
    
    doc.add_paragraph('PT: Implementação de autorização em rotas:')
    doc.add_paragraph('EN: Authorization implementation in routes:')
    
    add_code_block(doc, '''
router.post('/tables/:id/free', 
  authenticateToken,
  canFreeTable,
  freeTable
);
    ''')
    
    doc.add_page_break()
    
    # ======================================
    # 9. TESTES
    # ======================================
    add_heading_custom(doc, '9. Testes / Testing', 1)
    
    add_heading_custom(doc, '9.1 Cenários de Teste / Test Scenarios', 2)
    
    test_scenarios = [
        ('Teste 1', 'Test 1', 'Transição automática free → occupied', 'Automatic transition free → occupied'),
        ('Teste 2', 'Test 2', 'Liberação manual de mesa', 'Manual table release'),
        ('Teste 3', 'Test 3', 'Autorização por role', 'Authorization by role'),
        ('Teste 4', 'Test 4', 'Bloqueio de pedidos em mesa fechada', 'Blocking orders on closed table'),
        ('Teste 5', 'Test 5', 'Histórico de sessões', 'Session history')
    ]
    
    for num_pt, num_en, desc_pt, desc_en in test_scenarios:
        doc.add_paragraph(f'{num_pt} / {num_en}:', style='Heading 3')
        doc.add_paragraph(f'🇵🇹 {desc_pt}')
        doc.add_paragraph(f'🇬🇧 {desc_en}')
    
    doc.add_page_break()
    
    # ======================================
    # 10. INTERNACIONALIZAÇÃO
    # ======================================
    add_heading_custom(doc, '10. Internacionalização / Internationalization', 1)
    
    doc.add_paragraph('PT: O sistema suporta 4 idiomas com 33 novas chaves de tradução.')
    doc.add_paragraph('EN: The system supports 4 languages with 33 new translation keys.')
    
    doc.add_paragraph('Idiomas / Languages:', style='Heading 3')
    doc.add_paragraph('• Português (PT)')
    doc.add_paragraph('• English (EN)')
    doc.add_paragraph('• Español (ES)')
    doc.add_paragraph('• Français (FR)')
    
    doc.add_paragraph('Principais Chaves / Main Keys:', style='Heading 3')
    
    i18n_keys = [
        'view_orders', 'current_session', 'session_duration',
        'free_table', 'session_orders', 'order_count',
        'total_revenue', 'session_history'
    ]
    
    for key in i18n_keys:
        doc.add_paragraph(f'• {key}', style='List Bullet')
    
    doc.add_page_break()
    
    # ======================================
    # 11. CONCLUSÃO
    # ======================================
    add_heading_custom(doc, '11. Conclusão / Conclusion', 1)
    
    doc.add_paragraph(
        'PT: O Sistema de Gestão de Estado de Mesas foi implementado com sucesso, '
        'oferecendo uma solução robusta e completa para gerenciamento de mesas em restaurantes. '
        'O sistema inclui transições automáticas, controles de autorização, '
        'rastreamento de receita, e suporte multilíngue.'
    )
    
    doc.add_paragraph(
        'EN: The Table State Management System has been successfully implemented, '
        'providing a robust and complete solution for table management in restaurants. '
        'The system includes automatic transitions, authorization controls, '
        'revenue tracking, and multilingual support.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph('Status da Implementação / Implementation Status:', style='Heading 2')
    doc.add_paragraph('✅ Backend: 100% completo / 100% complete')
    doc.add_paragraph('✅ Frontend Admin: 100% completo / 100% complete')
    doc.add_paragraph('✅ i18n: 4 idiomas / 4 languages')
    doc.add_paragraph('✅ Testes: Prontos para execução / Ready for execution')
    doc.add_paragraph('✅ Documentação: Completa / Complete')
    
    # Salvar documento
    output_path = r'C:\Users\mpatricio\.gemini\antigravity\brain\d2e8fe11-6d00-43db-ad47-d4070f0fc7fc'
    if not os.path.exists(output_path):
        os.makedirs(output_path)
    
    file_path = os.path.join(output_path, 'Sistema_Gestao_Mesas_Technical_Documentation.docx')
    doc.save(file_path)
    
    print('Documento criado com sucesso!')
    print(f'Localizacao: {file_path}')
    return file_path

if __name__ == '__main__':
    try:
        path = create_technical_document()
        print(f'\nDocumento tecnico bilingue gerado!')
        print(f'Caminho: {path}')
    except Exception as e:
        print(f'Erro ao criar documento: {e}')
        import traceback
        traceback.print_exc()
